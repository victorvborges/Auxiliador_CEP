from flask import Flask, render_template, request, redirect, url_for, session, send_file
from agent import run_chapter_crew
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn

app = Flask(__name__)
app.secret_key = 'chave-secreta-super-segura'

projetos_em_memoria = {}

CAPITULOS_COM_HIPOTESE = [
    'Introdução / Referencial Teórico',
    'Problematização',
    'Justificativa',
    'Objeto de estudo',
    'Hipótese',
    'Objetivos',
    'Metodologia',
    'Aspectos éticos',
    'Análise e Interpretação dos Dados',
    'Cronograma',
    'Orçamento',
    'Referências',
    'Anexos/Apêndices'
]

CAPITULOS_SEM_HIPOTESE = [
    'Introdução / Referencial Teórico',
    'Problematização',
    'Justificativa',
    'Objeto de estudo',
    'Objetivos',
    'Metodologia',
    'Aspectos éticos',
    'Análise e Interpretação dos Dados',
    'Cronograma',
    'Orçamento',
    'Referências',
    'Anexos/Apêndices'
]

def get_projeto():
    projeto_id = session.get('projeto_id')
    return projetos_em_memoria.get(projeto_id)

@app.route('/')
def home():
    projeto = get_projeto()

    if not projeto:
        return render_template('index.html')

    conteudos = projeto.get('conteudos', [])
    capitulo_atual = projeto.get('capitulo_atual', 0)

    lista_capitulos = projeto['capitulos']
    projeto_finalizado = capitulo_atual >= len(lista_capitulos)

    return render_template('index.html',
                           conteudos=conteudos,
                           projeto_finalizado=projeto_finalizado)

@app.route('/start', methods=['POST'])
def start():
    projeto_id = str(len(projetos_em_memoria) + 1)

    incluir_hipotese = request.form.get('incluir_hipotese', 'sim')
    lista_capitulos = CAPITULOS_COM_HIPOTESE if incluir_hipotese == 'sim' else CAPITULOS_SEM_HIPOTESE

    projeto = {
        'titulo': request.form['titulo'],
        'pesquisador': request.form['pesquisador'],
        'orientador': request.form['orientador'],
        'coorientador': request.form.get('coorientador'),
        'cidade': request.form['cidade'],
        'ano': request.form['ano'],
        'incluir_hipotese': incluir_hipotese,
        'capitulos': lista_capitulos,
        'conteudos': [],
        'capitulo_atual': 0
    }

    projetos_em_memoria[projeto_id] = projeto
    session['projeto_id'] = projeto_id

    return redirect(url_for('generate_first_chapter'))

@app.route('/generate_first_chapter')
def generate_first_chapter():
    projeto = get_projeto()
    if not projeto:
        return redirect(url_for('home'))

    return gerar_proximo_capitulo(projeto)

@app.route('/next_chapter', methods=['POST'])
def next_chapter():
    projeto = get_projeto()
    if not projeto:
        return redirect(url_for('home'))

    conteudos = projeto['conteudos']
    capitulo_atual = projeto['capitulo_atual']

    texto_editado = request.form.get('texto_editado')

    if texto_editado:
        if capitulo_atual > 0 and (capitulo_atual - 1) < len(conteudos):
            conteudos[capitulo_atual - 1]['texto'] = texto_editado
        else:
            print(f"⚠️ Tentou salvar revisão inválida. capitulo_atual={capitulo_atual}, len={len(conteudos)}")

    return gerar_proximo_capitulo(projeto)

def gerar_proximo_capitulo(projeto):
    conteudos = projeto['conteudos']
    capitulo_atual = projeto['capitulo_atual']
    lista_capitulos = projeto['capitulos']

    if capitulo_atual >= len(lista_capitulos):
        print("✅ Todos os capítulos foram gerados!")
        return redirect(url_for('home'))

    proximo_capitulo = lista_capitulos[capitulo_atual]
    tema = projeto['titulo']
    contexto_anterior = "\n\n".join([c['texto'] for c in conteudos])

    texto_gerado = run_chapter_crew(proximo_capitulo, tema, contexto_anterior)
    print(f"✅ Capítulo gerado: {proximo_capitulo}")

    conteudos.append({
        'capitulo': proximo_capitulo,
        'texto': texto_gerado
    })

    projeto['capitulo_atual'] += 1

    print(f"📌 Estado após geração: capitulo_atual={projeto['capitulo_atual']}, len_conteudos={len(conteudos)}")

    return redirect(url_for('next_chapter_form'))

@app.route('/next_chapter_form', methods=['GET'])
def next_chapter_form():
    projeto = get_projeto()

    if not projeto:
        return redirect(url_for('home'))

    conteudos = projeto['conteudos']
    capitulo_atual = projeto['capitulo_atual']

    if capitulo_atual == 0 or len(conteudos) == 0:
        return redirect(url_for('home'))

    if capitulo_atual > len(conteudos):
        print("✅ Projeto finalizado!")
        return redirect(url_for('home'))

    capitulo_revisar = conteudos[capitulo_atual - 1]

    return render_template('chapter.html',
                           capitulo=capitulo_revisar['capitulo'],
                           resultado=capitulo_revisar['texto'])

@app.route('/reiniciar', methods=['POST'])
def reiniciar():
    projeto_id = session.get('projeto_id')
    if projeto_id and projeto_id in projetos_em_memoria:
        del projetos_em_memoria[projeto_id]

    session.clear()
    return redirect(url_for('home'))

@app.route('/export_docx')
def export_docx():
    projeto = get_projeto()
    if not projeto:
        return redirect(url_for('home'))

    conteudos = projeto['conteudos']
    titulo = projeto['titulo']
    pesquisador = projeto['pesquisador']
    orientador = projeto['orientador']
    coorientador = projeto['coorientador']
    cidade = projeto['cidade']
    ano = projeto['ano']

    doc = Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(85.05)
        section.bottom_margin = Pt(56.7)
        section.left_margin = Pt(85.05)
        section.right_margin = Pt(56.7)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Capa
    for text in [
        'UNIVERSIDADE FEDERAL DE PERNAMBUCO',
        'NOME DO CENTRO/NÚCLEO DA UFPE',
        'NOME DO CURSO',
        titulo.upper(),
        cidade.upper(),
        ano
    ]:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    doc.add_paragraph(pesquisador).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(titulo.upper()).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Projeto apresentado ao Comitê de Ética em Pesquisa.")
    doc.add_paragraph(f"Orientador(a): {orientador}")
    if coorientador:
        doc.add_paragraph(f"Coorientador(a): {coorientador}")
    doc.add_paragraph(f"{cidade} - {ano}")

    doc.add_page_break()

    doc.add_heading('SUMÁRIO', level=1)
    for idx, capitulo in enumerate(conteudos, start=1):
        doc.add_paragraph(f"{idx}. {capitulo['capitulo']} ........................................................... {idx + 1}")
    doc.add_page_break()

    for capitulo in conteudos:
        heading = doc.add_heading(capitulo['capitulo'], level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(capitulo['texto'])
        doc.add_page_break()

    output_path = 'Projeto_Pesquisa_CEP.docx'
    doc.save(output_path)

    return send_file(output_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
